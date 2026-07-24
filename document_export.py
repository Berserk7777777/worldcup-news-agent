import html
import re
from io import BytesIO
from pathlib import Path

from news_images import image_caption, images_by_placement


INK = "18352B"
MUTED = "68776F"


def _clean_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"^#{1,6}\s*", "", text.strip())
    return text.replace("**", "").replace("__", "").replace("`", "")


def _article_body(article: str) -> str:
    return re.split(r"\n\s*来源：\s*\n", article, maxsplit=1)[0].strip()


def _content(result: dict) -> dict:
    reviewer = result.get("reviewer_result", {})
    writer = result.get("writer_result", {})
    user_input = result.get("user_input", {})
    title = (
        reviewer.get("final_title")
        or writer.get("title")
        or user_input.get("topic")
        or "世界杯新闻"
    )
    article = reviewer.get("final_article") or writer.get("full_article", "")
    paragraphs = [
        _clean_markdown(item)
        for item in re.split(r"\n{2,}", _article_body(article))
        if item.strip()
    ]
    if paragraphs and re.sub(r"\s+", "", paragraphs[0]) == re.sub(
        r"\s+", "", _clean_markdown(title)
    ):
        paragraphs = paragraphs[1:]
    return {
        "title": _clean_markdown(title),
        "label": reviewer.get("final_article_label", ""),
        "created_at": result.get("created_at", "未知"),
        "news_type": user_input.get("news_type", "未指定"),
        "paragraphs": paragraphs,
        "sources": result.get("sources", []),
    }


def _set_run_font(run, name: str, size, color=None, bold=None) -> None:
    from docx.oxml.ns import qn

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    if not url:
        paragraph.add_run("未提供链接")
        return
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F6B4F")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    for child in hyperlink.iter(qn("w:r")):
        child_properties = child.find(qn("w:rPr"))
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(int(Pt(9).pt * 2)))
        child_properties.append(size)


def build_docx_bytes(result: dict, run_dir: Path) -> bytes:
    from PIL import Image as PillowImage
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    content = _content(result)
    placed_images = images_by_placement(result, run_dir)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    _set_run_font(
        kicker.add_run("2026 WORLD CUP NEWSROOM"),
        "Microsoft YaHei",
        Pt(9),
        RGBColor.from_string(MUTED),
        True,
    )

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    _set_run_font(
        title.add_run(content["title"]),
        "Microsoft YaHei",
        Pt(24),
        RGBColor.from_string(INK),
        True,
    )

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    _set_run_font(
        meta.add_run(
            f"生成时间：{content['created_at']}    新闻类型：{content['news_type']}"
        ),
        "Microsoft YaHei",
        Pt(9),
        RGBColor.from_string(MUTED),
    )

    def add_image(record: dict) -> None:
        path = record["path"]
        with PillowImage.open(path) as image_file:
            width, height = image_file.size
        if width / height >= 6.5 / 7:
            display_width = 6.5
            display_height = display_width * height / width
        else:
            display_height = 7
            display_width = display_height * width / height
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(
            str(path),
            width=Inches(display_width),
            height=Inches(display_height),
        )
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(12)
        _set_run_font(
            caption_paragraph.add_run(image_caption(record)),
            "Microsoft YaHei",
            Pt(9),
            RGBColor.from_string(MUTED),
        )

    for record in placed_images.get("cover", []):
        add_image(record)

    for index, text in enumerate(content["paragraphs"]):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(22)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.25
        _set_run_font(
            paragraph.add_run(text),
            "Microsoft YaHei",
            Pt(11),
            RGBColor.from_string("202B27"),
        )
        if index == 0:
            for record in placed_images.get("after_lead", []):
                add_image(record)
        if index == 1:
            for record in placed_images.get("after_paragraph_2", []):
                add_image(record)

    if not content["paragraphs"]:
        for record in placed_images.get("after_lead", []):
            add_image(record)
    if len(content["paragraphs"]) < 2:
        for record in placed_images.get("after_paragraph_2", []):
            add_image(record)
    for record in placed_images.get("gallery", []):
        add_image(record)

    if content["sources"]:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(8)
        _set_run_font(
            heading.add_run("来源"),
            "Microsoft YaHei",
            Pt(16),
            RGBColor.from_string(INK),
            True,
        )
        for index, source in enumerate(content["sources"], 1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            date = source.get("published_at") or "未标注日期"
            text = (
                f"[{index}] {source.get('source_name', '未知来源')}，"
                f"《{source.get('document_title', '未命名文章')}》，{date}，"
            )
            _set_run_font(
                paragraph.add_run(text),
                "Microsoft YaHei",
                Pt(9),
                RGBColor.from_string("404B46"),
            )
            _add_hyperlink(paragraph, "原始链接", source.get("source_url", ""))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        footer.add_run("2026 World Cup Newsroom · AI辅助新闻创作"),
        "Microsoft YaHei",
        Pt(8),
        RGBColor.from_string(MUTED),
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf_bytes(result: dict, run_dir: Path) -> bytes:
    from PIL import Image as PillowImage
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        Image,
        KeepTogether,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    content = _content(result)
    placed_images = images_by_placement(result, run_dir)
    output = BytesIO()
    windows_fonts = Path("C:/Windows/Fonts")
    if (windows_fonts / "msyh.ttc").exists():
        pdfmetrics.registerFont(TTFont("NewsFont", windows_fonts / "msyh.ttc"))
        pdfmetrics.registerFont(TTFont("NewsFontBold", windows_fonts / "msyhbd.ttc"))
        body_font = "NewsFont"
        title_font = "NewsFontBold"
    else:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        body_font = title_font = "STSong-Light"
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=content["title"],
        author="2026 World Cup Newsroom",
    )
    styles = {
        "kicker": ParagraphStyle(
            "Kicker",
            fontName=body_font,
            fontSize=9,
            leading=12,
            textColor=HexColor("#68776F"),
            spaceAfter=10,
        ),
        "title": ParagraphStyle(
            "NewsTitle",
            fontName=title_font,
            fontSize=24,
            leading=32,
            textColor=HexColor("#18352B"),
            spaceAfter=9,
        ),
        "meta": ParagraphStyle(
            "Meta",
            fontName=body_font,
            fontSize=9,
            leading=13,
            textColor=HexColor("#68776F"),
            spaceAfter=18,
        ),
        "body": ParagraphStyle(
            "Body",
            fontName=body_font,
            fontSize=11,
            leading=18,
            firstLineIndent=22,
            alignment=TA_LEFT,
            wordWrap="CJK",
            textColor=HexColor("#202B27"),
            spaceAfter=8,
        ),
        "caption": ParagraphStyle(
            "Caption",
            fontName=body_font,
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
            textColor=HexColor("#68776F"),
            spaceAfter=12,
        ),
        "heading": ParagraphStyle(
            "Heading",
            fontName=title_font,
            fontSize=16,
            leading=20,
            textColor=HexColor("#18352B"),
            spaceBefore=12,
            spaceAfter=8,
        ),
        "source": ParagraphStyle(
            "Source",
            fontName=body_font,
            fontSize=9,
            leading=14,
            wordWrap="CJK",
            textColor=HexColor("#404B46"),
            spaceAfter=6,
        ),
    }

    story = [
        Paragraph("2026 WORLD CUP NEWSROOM", styles["kicker"]),
        Paragraph(html.escape(content["title"]), styles["title"]),
        Paragraph(
            html.escape(
                f"生成时间：{content['created_at']}    新闻类型：{content['news_type']}"
            ),
            styles["meta"],
        ),
    ]
    max_width = 6.5 * inch
    max_height = 4.5 * inch

    def image_flowables(record: dict) -> KeepTogether:
        path = record["path"]
        with PillowImage.open(path) as image_file:
            width, height = image_file.size
        scale = min(max_width / width, max_height / height)
        picture = Image(str(path), width=width * scale, height=height * scale)
        picture.hAlign = "CENTER"
        return KeepTogether(
            [
                picture,
                Spacer(1, 6),
                Paragraph(html.escape(image_caption(record)), styles["caption"]),
            ]
        )

    story.extend(image_flowables(record) for record in placed_images.get("cover", []))
    for index, text in enumerate(content["paragraphs"]):
        story.append(Paragraph(html.escape(text), styles["body"]))
        if index == 0:
            story.extend(
                image_flowables(record)
                for record in placed_images.get("after_lead", [])
            )
        if index == 1:
            story.extend(
                image_flowables(record)
                for record in placed_images.get("after_paragraph_2", [])
            )
    if not content["paragraphs"]:
        story.extend(
            image_flowables(record)
            for record in placed_images.get("after_lead", [])
        )
    if len(content["paragraphs"]) < 2:
        story.extend(
            image_flowables(record)
            for record in placed_images.get("after_paragraph_2", [])
        )
    story.extend(image_flowables(record) for record in placed_images.get("gallery", []))
    if content["sources"]:
        story.append(Paragraph("来源", styles["heading"]))
        for index, source in enumerate(content["sources"], 1):
            date = source.get("published_at") or "未标注日期"
            source_text = (
                f"[{index}] {source.get('source_name', '未知来源')}，"
                f"《{source.get('document_title', '未命名文章')}》，{date}，"
                f'<link href="{html.escape(source.get("source_url", ""), quote=True)}" '
                'color="#0F6B4F">原始链接</link>'
            )
            story.append(Paragraph(source_text, styles["source"]))

    def footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(body_font, 8)
        canvas.setFillColor(HexColor("#68776F"))
        canvas.drawCentredString(
            letter[0] / 2,
            0.45 * inch,
            f"2026 World Cup Newsroom · 第 {doc.page} 页",
        )
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()
